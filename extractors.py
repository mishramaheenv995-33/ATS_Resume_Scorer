import pdfplumber
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from PIL import Image
import os

def extract_text_pdf(file_path):
    """Extract text from PDF using pdfplumber"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                print(f"[INFO] Page {page_num}: Extracted {len(page_text) if page_text else 0} characters")
    except Exception as e:
        print(f"[ERROR] PDF extraction failed: {e}")
    return text

def extract_text_image_pdf(file_path):
    """Extract text from image-based PDF using OCR"""
    text = ""
    try:
        images = convert_from_path(file_path, dpi=300)   
        print(f"[INFO] Converting {len(images)} pages to images for OCR...")
        
        for i, image in enumerate(images, 1):
             
            image = image.convert('L')   
            page_text = pytesseract.image_to_string(image, config='--psm 6')
            if page_text.strip():
                text += page_text + "\n"
            print(f"[INFO] OCR Page {i}: Extracted {len(page_text)} characters")
            
    except Exception as e:
        print(f"[ERROR] OCR extraction failed: {e}")
    return text

def extract_text_image(file_path):
    """Extract text from image files using OCR"""
    text = ""
    try:
        image = Image.open(file_path)
        
        if image.mode != 'L':
            image = image.convert('L')
        
         
        text = pytesseract.image_to_string(image, config='--psm 6')
        print(f"[INFO] Image OCR: Extracted {len(text)} characters")
        
    except Exception as e:
        print(f"[ERROR] Image OCR failed: {e}")
    return text

def extract_text_docx(file_path):
    """Extract text from Word document"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
         
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        print(f"[INFO] DOCX: Extracted {len(text)} characters")
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed: {e}")
    return text

def extract_text_resume(file_path):
    """
    Main function to extract text from resume files
    Supports PDF, DOCX, and common image formats
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    if ext == ".pdf":
        print("[INFO] Processing PDF file...")
        text = extract_text_pdf(file_path)
        
         
        if not text.strip():
            print("[INFO] No text found with pdfplumber, trying OCR...")
            text = extract_text_image_pdf(file_path)
        
    elif ext in [".docx", ".doc"]:
        print("[INFO] Processing Word document...")
        text = extract_text_docx(file_path)
        
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".gif"]:
        print("[INFO] Processing image file with OCR...")
        text = extract_text_image(file_path)
        
    else:
        raise ValueError(f"Unsupported file format: {ext}. "
                        "Supported formats: PDF, DOCX, JPG, PNG, BMP, TIFF")
    
    if not text.strip():
        print("[WARNING] No text could be extracted from the file")
    else:
        print(f"[SUCCESS] Total text extracted: {len(text)} characters")
    
    return text
